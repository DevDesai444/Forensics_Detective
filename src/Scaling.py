# import os
# import random
# import math
# import shutil
# from pathlib import Path
# import PyPDF2
# import wikipedia
# import matplotlib.pyplot as plt
# from wikipedia import page, search
# from wikipedia.exceptions import DisambiguationError, PageError
# import pandas as pd
# import pandas as pd
# from docx import Document
# from docx.shared import Inches
# # Uncomment if using docx2pdf for PDF conversion
# # from docx2pdf import convert

# # Configuration
# MOTHER_DIR = Path("/Users/DEVDESAI1/Desktop/University_at_Buffalo/EAS510/Forensics_Detective/src/Mother_Pdfs")
# OUTPUT_DIR = Path("scaled_pdfs")
# NUM_DOCS = 10000
# TOTAL_PAGES = 10000
# WIKI_TOPICS = [
#     "Machine learning", "Computer vision", "Deep learning", "Neural network",
#     "Data science", "Image processing", "Python (programming language)",
#     "Artificial intelligence", "PDF", "Document classification"
# ]

# # Ensure output folders
# OUTPUT_DIR.mkdir(exist_ok=True)
# (PDF_DIR := OUTPUT_DIR / "pdfs").mkdir(exist_ok=True)
# (DOCX_DIR := OUTPUT_DIR / "docx").mkdir(exist_ok=True)

# # Precompute how many pages per doc (simple uniform for total pages)
# pages_per_doc = TOTAL_PAGES // NUM_DOCS

# # def fetch_wiki_content(title):
# #     """Fetch summary, tables and an image URL from Wikipedia."""
# #     try:
# #         page = wikipedia.page(title)
# #     except wikipedia.DisambiguationError as e:
# #         page = wikipedia.page(e.options[0])
# #     txt = page.summary
# #     # Simple table: take first HTML table via pandas
# #     try:
# #         tables = pd.read_html(page.url)
# #         table = tables[0]
# #     except:
# #         table = None
# #     # Wikipedia images list
# #     img_url = page.images[0] if page.images else None
# #     return txt, table, img_url

# def fetch_wiki_content(title):
#     """Fetch summary, first table, and first image URL from Wikipedia."""
#     # Normalize title
#     clean_title = title.strip().replace(";", "").title()
    
#     # Attempt to load the page, with fallbacks for disambiguation and missing pages
#     try:
#         wiki_page = page(clean_title)
#     except DisambiguationError as e:
#         wiki_page = page(e.options[0])
#     except PageError:
#         results = search(clean_title)
#         if not results:
#             return None, None, None
#         try:
#             wiki_page = page(results[0])
#         except (DisambiguationError, PageError):
#             return None, None, None
    
#     # Extract summary
#     txt = wiki_page.summary
    
#     # Extract first HTML table, if any
#     try:
#         tables = pd.read_html(wiki_page.url)
#         table = tables[0]
#     except Exception:
#         table = None
    
#     # Extract first image URL, if available
#     img_url = wiki_page.images[0] if wiki_page.images else None
    
#     return txt, table, img_url

# def make_chart(doc, data_frame):
#     """Embed a chart for a pandas DataFrame into the doc."""
#     fig, ax = plt.subplots(figsize=(4,3))
#     data_frame.head(5).plot(kind="bar", ax=ax)
#     img_path = "temp_chart.png"
#     plt.tight_layout()
#     fig.savefig(img_path)
#     plt.close(fig)
#     doc.add_picture(img_path, width=Inches(3.5))
#     os.remove(img_path)

# def process_mother_pdf(path):
#     """Extract text pages and images from mother PDF."""
#     reader = PyPDF2.PdfReader(str(path))
#     texts = []
#     for page in reader.pages:
#         texts.append(page.extract_text())
#     # For brevity, skipping image extraction; you could use pdf2image
#     return texts

# # Load all mother texts in memory
# mother_texts = []
# for pdf_file in MOTHER_DIR.glob("*.pdf"):
#     mother_texts.extend(process_mother_pdf(pdf_file))

# for doc_idx in range(1, NUM_DOCS+1):
#     doc = Document()
#     # Randomly choose a Wikipedia topic each document
#     title = random.choice(WIKI_TOPICS)
#     summary, table, img_url = fetch_wiki_content(title)
    
#     # Build pages
#     for p in range(pages_per_doc):
#         # Insert original text snippet
#         snippet = random.choice(mother_texts[:400])[:500]  # first 500 chars
#         doc.add_paragraph(snippet)
        
#         # Occasionally embed wiki summary
#         # if random.random() < 0.3:
#         if summary and random.random() < 0.3:
#             doc.add_heading(f"Wikipedia Excerpt: {title}", level=2)
#             doc.add_paragraph(summary[:1000])
        
#         # Embed table/table chart per complexity requirement
#         if table is not None and random.random() < 0.25:
#             doc.add_paragraph("Data Table Example:")
#             # Insert table
#             tbl = doc.add_table(rows=1, cols=len(table.columns))
#             hdr_cells = tbl.rows[0].cells
#             for i,col in enumerate(table.columns.tolist()):
#                 hdr_cells[i].text = str(col)
#             for _, row in table.head(3).iterrows():
#                 row_cells = tbl.add_row().cells
#                 for i,val in enumerate(row):
#                     row_cells[i].text = str(val)
#             # Also add a chart
#             make_chart(doc, table)
        
#         # Embed image if URL available
#         if img_url and random.random() < 0.3:
#             # Download and embed image
#             import requests
#             img_resp = requests.get(img_url, stream=True)
#             if img_resp.status_code == 200:
#                 img_file = f"temp_img_{doc_idx}.png"
#                 with open(img_file, "wb") as f:
#                     shutil.copyfileobj(img_resp.raw, f)
#                 doc.add_picture(img_file, width=Inches(3.0))
#                 os.remove(img_file)
        
#         # Add a page break
#         if p < pages_per_doc-1:
#             doc.add_page_break()
    
#     # Save DOCX
#     docx_path = DOCX_DIR / f"doc_{doc_idx:05d}.docx"
#     doc.save(docx_path)
    
#     # Convert to PDF (choose one method below)
#     # Method A: docx2pdf (Windows/Mac)
#     # convert(str(docx_path), str(PDF_DIR / f"doc_{doc_idx:05d}.pdf"))
    
#     # Method B: LibreOffice headless
    
    
    
#     # os.system(f'soffice --headless --convert-to pdf "{docx_path}" --outdir "{PDF_DIR}"')
#     # if doc_idx % 100 == 0:
#     #     print(f"Generated {doc_idx} / {NUM_DOCS}")

# print("All documents generated.")


# import os, random, shutil, requests
# from pathlib import Path
# import PyPDF2
# import wikipedia
# from wikipedia.exceptions import DisambiguationError, PageError
# import pandas as pd
# import matplotlib.pyplot as plt
# from docx import Document
# from docx.shared import Inches
# from sympy import latex, symbols, Eq  # for equations
# from docx2pdf import convert  # optional for PDF
# # Ensure LibreOffice headless export for .pages

# # CONFIG
# MOTHER = Path("/Users/DEVDESAI1/Desktop/University_at_Buffalo/EAS510/Forensics_Detective/src/Mother_Pdfs")
# OUT_word = Path("scaled_word")
# OUT_pages = Path("scaled_pdf")
# TOPICS = [f.stem for f in MOTHER.glob("*.pdf")]
# # WIKI = ["Machine learning","Computer vision","Deep learning","Neural network",
# #         "Data science","Image processing","Python (programming language)",
# #         "Artificial intelligence","PDF","Document classification"]

# # HELPERS
# def fetch_page(title):
#     t=title.strip().title()
#     try: p=wikipedia.page(t)
#     except DisambiguationError as e: p=wikipedia.page(e.options[0])
#     except PageError:
#         r=wikipedia.search(t)
#         if not r: return None
#         p=wikipedia.page(r[0])
#     return p

# def insert_wiki_image(doc,p,prob=0.3):
#     if p and p.images and random()<prob:
#         url=p.images[0]; r=requests.get(url,stream=True)
#         if r.ok:
#             fn="tmp.png"
#             with open(fn,"wb") as f: shutil.copyfileobj(r.raw,f)
#             doc.add_picture(fn,width=Inches(3)); os.remove(fn)

# def insert_wiki_chart(doc,p,prob=0.3):
#     if p and random()<prob:
#         try:
#             df=pd.read_html(p.url)[0].head(5)
#             fig,ax=plt.subplots()
#             df.plot.bar(ax=ax); fig.savefig("tmpc.png"); plt.close(fig)
#             doc.add_picture("tmpc.png",width=Inches(3)); os.remove("tmpc.png")
#         except: pass

# def insert_wiki_diagram(doc,p,prob=0.3):
#     if p and random()<prob:
#         # use matplotlib flowchart placeholder
#         fig,ax=plt.subplots()
#         ax.text(0.5,0.5,"Diagram",ha="center")
#         fig.savefig("tmpd.png"); plt.close(fig)
#         doc.add_picture("tmpd.png",width=Inches(3)); os.remove("tmpd.png")

# def insert_table(doc,p,prob=0.3):
#     if p and random()<prob:
#         try:
#             df=pd.read_html(p.url)[0].sample(3)
#             tbl=doc.add_table(1,len(df.columns))
#             for i,c in enumerate(df.columns): tbl.rows[0].cells[i].text=str(c)
#             for _,r in df.iterrows():
#                 row=tbl.add_row().cells
#                 for i,v in enumerate(r): row[i].text=str(v)
#         except: pass

# def insert_equation(doc,prob=0.3):
#     if random()<prob:
#         x,y=symbols('x y')
#         eq=Eq(x**2+y**2,1)
#         doc.add_paragraph(f"Equation: ${latex(eq)}$")

# def insert_headings_lists(doc):
#     doc.add_heading("Section 1",level=1)
#     doc.add_paragraph("• Bullet A\n• Bullet B")
#     doc.add_paragraph("1. First\n2. Second")

# def insert_citations_footnotes(doc):
#     doc.add_paragraph("This method was shown effective in prior work[1].")
#     # footnotes not supported; simulate as bracket

# def expand_content(doc,pages_target):
#     # duplicate content until pages >= target
#     while len(doc.element.body.findall('.//w:sectPr'))<pages_target:
#         doc.add_paragraph("Additional filler content to reach page target.")

# def process_pdf(src):
#     txts=[]
#     r=PyPDF2.PdfReader(str(src))
#     for pg in r.pages: txts.append(pg.extract_text())
#     return txts

# for pdf in MOTHER.glob("*.pdf"):
#     base=pdf.stem
#     texts=process_pdf(pdf)
#     # PHASE 1: 6 variants
#     mods = [
#         ("img", insert_wiki_image),
#         ("chart", insert_wiki_chart),
#         ("diag", insert_wiki_diagram),
#         ("img_chart", lambda d,p: (insert_wiki_image(d,p),insert_wiki_chart(d,p))),
#         ("img_diag", lambda d,p: (insert_wiki_image(d,p),insert_wiki_diagram(d,p))),
#         ("chart_diag", lambda d,p: (insert_wiki_chart(d,p),insert_wiki_diagram(d,p))),
#     ]
#     # PHASE 2: 7 variants with tables + (diagram/chart/images)
#     tbl_mods = [
#        ("tbl_diag", insert_wiki_diagram),
#        ("tbl_chart", insert_wiki_chart),
#        ("tbl_img", insert_wiki_image),
#        ("tbl", None)
#     ]
#     # generate
#     for name,fn in mods:
#         doc=Document()
#         # p=fetch_page(random.choice(WIKI))
#         topic = base.replace("_", " ").title()
#         p = fetch_page(topic)
#         if p is None:
#            # fallback: search by stem
#            p = fetch_page(base)
#         doc.add_paragraph(random.choice(texts) or "")
#         fn(doc,p) if fn else None
#         expand_content(doc,random.randint(1,5))
#         out=OUT_word/f"{base}_{name}.docx"; doc.save(out)
#         os.system(f'soffice --headless --convert-to pages "{out}" --outdir "{OUT_pages}"')
#     for name,fn in tbl_mods:
#         for i in range(2 if name!="tbl" else 1):
#             doc=Document(); 
#             # p=fetch_page(random.choice(WIKI))
#             topic = base.replace("_", " ").title()
#             p = fetch_page(topic)
#             if p is None:
#             # fallback: search by stem
#                 p = fetch_page(base)
#             insert_table(doc,p)
#             if fn: fn(doc,p)
#             expand_content(doc,random.randint(1,5))
#             out=OUT_word/f"{base}_{name}{i}.docx"; doc.save(out)
#             os.system(f'soffice --headless --convert-to pages "{out}" --outdir "{OUT_pages}"')
#     # PHASE 3: 4 complex
#     for i in range(4):
#         doc=Document(); 
#         # p=fetch_page(random.choice(WIKI))
#         topic = base.replace("_", " ").title()
#         p = fetch_page(topic)
#         if p is None:
#            # fallback: search by stem
#            p = fetch_page(base)
#         insert_headings_lists(doc)
#         insert_equation(doc)
#         insert_citations_footnotes(doc)
#         expand_content(doc,random.randint(1,5))
#         out=OUT_word/f"{base}_complex{i}.docx"; doc.save(out)
#         os.system(f'soffice --headless --convert-to pages "{out}" --outdir "{OUT_pages}"')
# print("Completed scaling for all 300 PDFs.")



#!/usr/bin/env python3
"""
Batch-generate .docx and .pages variants of source PDFs with enriched content
from Wikipedia: embedded images, charts, diagrams, tables, lists, equations,
citations, and footnotes. 

Dependencies:
  - python-docx
  - wikipedia
  - requests
  - pillow
  - tqdm

Usage:
  python3 Scaling.py
"""

import os
import io
import random
import shutil
import requests
import wikipedia
from docx import Document
from docx.shared import Inches
from PIL import Image
from tqdm import tqdm

# Base directories
MOTHER = "/Users/DEVDESAI1/Desktop/University_at_Buffalo/EAS510/Forensics_Detective/src/Mother_Pdfs"
OUT_word = "scaled_word"
OUT_pages = "scaled_pdf"

# Ensure output subdirs exist
for base in (OUT_word, OUT_pages):
    for mode in [
        "images", "charts", "diagrams",
        "images_charts", "images_diagrams", "charts_diagrams",
        "tables_diagrams_1", "tables_charts_2", "tables_diagrams_2", "tables_only",
        "complex_1", "complex_2", "complex_3", "complex_4"
    ]:
        os.makedirs(os.path.join(base, mode), exist_ok=True)

def get_wikipedia_media(topic):
    """
    Fetch a random image/chart/diagram URL from the Wikipedia page for topic.
    """
    try:
        page = wikipedia.page(topic)
    except Exception:
        return None
    imgs = [u for u in page.images if u.lower().endswith((".png", ".jpg", ".svg"))]
    return random.choice(imgs) if imgs else None

# def download_image(url):
#     resp = requests.get(url)
#     resp.raise_for_status()
#     return Image.open(io.BytesIO(resp.content))
# def download_image(url):
#     headers = {
#         "User-Agent": "ScalingScript/1.0 (devdesaiyt@gmail.com)"
#     }
#     # If it’s an SVG, use the PNG export endpoint
#     if url.lower().endswith(".svg"):
#         url = url.replace(
#             "https://upload.wikimedia.org/wikipedia/commons/",
#             "https://upload.wikimedia.org/wikipedia/commons/thumb/"
#         ) + "/1024px-" + os.path.basename(url).replace(".svg", ".png")
#     resp = requests.get(url, headers=headers)
#     resp.raise_for_status()
#     return Image.open(io.BytesIO(resp.content))

# def download_image(url):
#     headers = {
#         "User-Agent": "ScalingScript/1.0 (devdesaiyt@gmail.com)"
#     }
#     # If it’s an SVG, rewrite to the Commons “thumb” PNG URL
#     if url.lower().endswith(".svg"):
#         # e.g. https://upload.wikimedia.org/wikipedia/commons/4/41/Global_thinking.svg
#         parts = url.split('/') # ['https://upload.wikimedia.org/wikipedia/en/9/96/Symbol_category_class.svg']
#         print(parts)
#         path = parts[-1]  # "4/41/Global_thinking.svg"
#         thumb_base = "https://upload.wikimedia.org/wikipedia/commons/thumb/" + path
#         filename = os.path.basename(path)  # "Global_thinking.svg"
#         png_name = "1024px-" + filename + ".png"
#         url = f"{thumb_base}/{png_name}"
#     resp = requests.get(url, headers=headers)
#     resp.raise_for_status()
#     return Image.open(io.BytesIO(resp.content))

# def download_image(url):
#     headers = {
#         "User-Agent": "ScalingScript/1.0 (your_email@domain.com)"
#     }
#     # If it’s an SVG, rewrite to the Commons “thumb” PNG URL
#     if url.lower().endswith(".svg"):
#         # e.g. https://upload.wikimedia.org/wikipedia/commons/4/41/Global_thinking.svg
#         parts = url.split("/wikipedia/commons/")
#         path = parts[1]  # "4/41/Global_thinking.svg"
#         thumb_base = "https://upload.wikimedia.org/wikipedia/commons/thumb/" + path
#         filename = os.path.basename(path)  # "Global_thinking.svg"
#         png_name = "1024px-" + filename + ".png"
#         url = f"{thumb_base}/{png_name}"
#     resp = requests.get(url, headers=headers)
#     resp.raise_for_status()
#     return Image.open(io.BytesIO(resp.content))

# def download_image(url):
#     headers = {
#         "User-Agent": "ScalingScript/1.0 (devdesaiyt@gmail.com)"
#     }
#     # Only rewrite if it’s an SVG in /wikipedia/commons/
#     if url.lower().endswith(".svg") and "/wikipedia/commons/" in url:
#         prefix, path = url.split("/wikipedia/commons/", 1)
#         thumb_base = prefix + "/wikipedia/commons/thumb/" + path
#         filename = os.path.basename(path)  # e.g. "Global_thinking.svg"
#         png_name = "1024px-" + filename + ".png"
#         url = f"{thumb_base}/{png_name}"
#     resp = requests.get(url, headers=headers)
#     resp.raise_for_status()
#     return Image.open(io.BytesIO(resp.content))

'''TRY 5'''
# WIKIPEDIA_API_URL = "https://en.wikipedia.org/w/api.php"

# def get_wikimedia_thumbnail_url(topic, width=512):
#     """
#     Query the MediaWiki API for a guaranteed thumbnail URL (PNG/JPG).
#     Returns None if no thumbnail is available.
#     """
#     params = {
#         "action": "query",
#         "titles": topic,
#         "prop": "pageimages",
#         "format": "json",
#         "pithumbsize": width
#     }
#     resp = requests.get(WIKIPEDIA_API_URL, params=params, timeout=10)
#     resp.raise_for_status()
#     data = resp.json().get("query", {}).get("pages", {})
#     for page in data.values():
#         thumb = page.get("thumbnail", {}).get("source")
#         if thumb:
#             return thumb
#     return None

# def download_image(topic):
#     """
#     Fetches a reliable raster image (PNG/JPG) for the given Wikipedia topic.
#     Uses the MediaWiki API to avoid SVG or forbidden‐content issues.
#     """
#     thumb_url = get_wikimedia_thumbnail_url(topic, width=1024)
#     if not thumb_url:
#         raise ValueError(f"No thumbnail found for topic: {topic}")
#     headers = {
#         "User-Agent": "ScalingScript/1.0 (devdesaiyt@gmail.com)"
#     }
#     # Optional: preflight HEAD to verify content type
#     head = requests.head(thumb_url, headers=headers, timeout=5)
#     head.raise_for_status()
#     content_type = head.headers.get("Content-Type", "")
#     if not content_type.startswith("image/"):
#         raise ValueError(f"Unexpected MIME type: {content_type}")
#     # Download and open
#     resp = requests.get(thumb_url, headers=headers, timeout=10)
#     resp.raise_for_status()
#     return Image.open(io.BytesIO(resp.content))

'''Try 7'''
WIKIPEDIA_API_URL = "https://en.wikipedia.org/w/api.php"
HEADERS = {"User-Agent": "ScalingScript/1.0 (devdesaiyt@gmail.com)"}

def resolve_wikipedia_title(query):
    return wikipedia.search(query, results=1)[0]

def get_wikimedia_thumbnail_url(query, width=512):
    title = resolve_wikipedia_title(query)
    params = {
        "action": "query",
        "titles": title,
        "prop": "pageimages",
        "format": "json",
        "pithumbsize": width
    }
    resp = requests.get(WIKIPEDIA_API_URL, params=params, headers=HEADERS, timeout=10)
    resp.raise_for_status()
    pages = resp.json()["query"]["pages"]
    for page in pages.values():
        thumb = page.get("thumbnail", {}).get("source")
        if thumb:
            return thumb
    # Fallback to first image URL
    return wikipedia.page(title).images[0]

def download_image(query):
    thumb_url = get_wikimedia_thumbnail_url(query, width=1024)
    resp = requests.get(thumb_url, headers=HEADERS, timeout=10)
    resp.raise_for_status()
    return Image.open(io.BytesIO(resp.content))

'''TRY 6'''
# UNSPLASH_ACCESS_KEY = None
# UNSPLASH_QUERIES    = ["nature","technology","city","science","art","animals"]
# IMAGES_CACHE_DIR    = "./_unsplash_images"
# os.makedirs(IMAGES_CACHE_DIR, exist_ok=True)

# def download_image(topic):
#     """
#     Fetch a random landscape image for the given topic from Unsplash,
#     save locally, and return the file path.
#     """
#     query = random.choice(UNSPLASH_QUERIES + [topic])
#     url   = f"https://api.unsplash.com/photos/random?query={query}" \
#             f"&client_id={UNSPLASH_ACCESS_KEY}&orientation=landscape"
#     resp  = requests.get(url, timeout=10)
#     resp.raise_for_status()
#     data  = resp.json()
#     img_url = data["urls"]["regular"]
#     img_resp = requests.get(img_url, timeout=10)
#     img_resp.raise_for_status()

#     img = Image.open(io.BytesIO(img_resp.content))
#     fname = f"{topic.replace(' ','_')}_{int(time.time()*1000)}.jpg"
#     path  = os.path.join(IMAGES_CACHE_DIR, fname)
#     img.save(path, format="JPEG")
#     return path

# def download_image(topic, width=None, height=None):
#     """
#     Fetch a random landscape image for the given topic from Unsplash
#     and return it as a PIL Image object, matching download_image’s output.
#     """
#     query = random.choice(UNSPLASH_QUERIES + [topic])
#     url = f"https://api.unsplash.com/photos/random?query={query}&client_id={UNSPLASH_ACCESS_KEY}&orientation=landscape"
#     resp = requests.get(url, timeout=10)
#     resp.raise_for_status()
#     data = resp.json()
#     img_url = data["urls"]["regular"]

#     img_resp = requests.get(img_url, timeout=10)
#     img_resp.raise_for_status()
#     img = PILImage.open(io.BytesIO(img_resp.content))

    # Optionally resize
    # if width and height:
    #     img = img.resize((width, height), PILImage.LANCZOS)

    # return img

'''Continue'''

def convert_to_pages(docx_path, pages_path):
    # Only run pandoc if it’s installed
    if shutil.which("pandoc"):
        os.system(f"pandoc {docx_path!r} -o {pages_path!r}")
    else:
        print("Warning: pandoc not found, skipping .pages conversion")

def extend_with_placeholder_text(doc, target_pages=3):
    """
    Add placeholder paragraphs until approx target_pages.
    """
    lorem = ("Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
             "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.")
    while len(doc.paragraphs) < target_pages * 3:
        doc.add_paragraph(lorem)

def create_variant_docx(topic, mode, flags):
    doc = Document()
    doc.add_heading(topic, level=0)

    # Headings
    if flags.get("headings"):
        doc.add_heading("Section A", level=1)
        doc.add_heading("Subsection A.1", level=2)

    # Media
    for media in ("images", "charts", "diagrams"):
        if flags.get(media):
            # url = get_wikipedia_media(topic)
            # url = get_wikimedia_thumbnail_url(topic)
            # if url:
            #     img = download_image(url)
            #     buf = io.BytesIO()
            #     img.save(buf, format="PNG")
            #     buf.seek(0)
            #     doc.add_picture(buf, width=Inches(3))
            
            img = download_image(topic)
            # Write it to a buffer and insert:
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            buf.seek(0)
            doc.add_picture(buf, width=Inches(3))


    # Tables
    if flags.get("tables"):
        tbl = doc.add_table(rows=1, cols=3)
        hdr = tbl.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Col 1", "Col 2", "Col 3"
        for _ in range(3):
            row = tbl.add_row().cells
            row[0].text = str(random.randint(1,100))
            row[1].text = str(random.randint(101,200))
            row[2].text = str(random.randint(201,300))

    # Lists
    if flags.get("lists"):
        doc.add_paragraph("Key points:", style="List Bullet")
        doc.add_paragraph("First item", style="List Bullet")
        doc.add_paragraph("Second item", style="List Number")

    # Equations
    if flags.get("equations"):
        p = doc.add_paragraph()
        p.add_run("Equation: ").bold = True
        p.add_run("E = mc²")

    # Citations and footnotes
    footnotes = []
    if flags.get("citations"):
        p = doc.add_paragraph("Reference")
        p.add_run(" (Wikipedia)")
    if flags.get("footnotes"):
        p = doc.add_paragraph("See note")
        p.add_run("[1]")
        footnotes.append("Generated footnote text.")
    if footnotes:
        doc.add_page_break()
        doc.add_heading("Footnotes", level=1)
        for i, note in enumerate(footnotes, 1):
            doc.add_paragraph(f"[{i}] {note}")

    # Extend to multiple pages
    extend_with_placeholder_text(doc, target_pages=random.randint(1,5))

    # Save .docx
    out_docx = os.path.join(OUT_word, mode, f"{topic}_{mode}.docx")
    doc.save(out_docx)

    # Convert to .pages via Pandoc
    out_pages = os.path.join(OUT_pages, mode, f"{topic}_{mode}.pages")
    # os.system(f"pandoc {out_docx} -o {out_pages}")
    out_docx = os.path.join(OUT_word, mode, f"{topic}_{mode}.docx")
    doc.save(out_docx)
    out_pages = os.path.join(OUT_pages, mode, f"{topic}_{mode}.pages")
    convert_to_pages(out_docx, out_pages)

def main():
    pdfs = [f for f in os.listdir(MOTHER) if f.lower().endswith(".pdf")]
    modes = {
        "images":            {"images":True},
        "charts":            {"charts":True},
        "diagrams":          {"diagrams":True},
        "images_charts":     {"images":True, "charts":True},
        "images_diagrams":   {"images":True, "diagrams":True},
        "charts_diagrams":   {"charts":True, "diagrams":True},
        "tables_diagrams_1": {"tables":True, "diagrams":True},
        "tables_charts_2":   {"tables":True, "charts":True},
        "tables_diagrams_2": {"tables":True, "diagrams":True},
        "tables_only":       {"tables":True},
        "complex_1":         {"headings":True, "lists":True, "equations":True, "citations":True, "footnotes":True},
        "complex_2":         {"headings":True, "lists":True, "equations":True, "citations":True, "footnotes":True},
        "complex_3":         {"headings":True, "lists":True, "equations":True, "citations":True, "footnotes":True},
        "complex_4":         {"headings":True, "lists":True, "equations":True, "citations":True, "footnotes":True},
    }

    for pdf in tqdm(pdfs, desc="Processing PDFs"):
        topic = os.path.splitext(pdf)[0]
        for mode, flags in modes.items():
            create_variant_docx(topic, mode, flags)

if __name__ == "__main__":
    main()